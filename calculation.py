import os
import random
from docx import Document
from docx.shared import Pt
import time


class Calculation:
    max = 5
    addAmount = 50
    minusAmount = 50
    addList = []
    minusList = []
    ITEM_LEN = 37

    def __init__(self, max=5, addAmount=50, minusAmount=50):
        self.max = max
        self.addAmount = addAmount
        self.minusAmount = minusAmount

    def add(self):
        candidates = set()
        for i in range(0, self.max+1):
            j = self.max - i
            for k in range(0, j+1):
                candidates.add(str(i)+' + '+str(k)+' =')
        #print(candidates)
        self.addList = random.sample(candidates, k=self.addAmount)

    def minus(self):
        candidates = set()
        for i in range(0, self.max+1):
            for j in range(0, i+1):
                candidates.add(str(i) + ' - ' + str(j) + ' =')
        #print(candidates)
        self.minusList = random.sample(candidates, k=self.minusAmount)

    def generate_doc(self):
        arr = self.addList + self.minusList
        arr = random.sample(arr, len(arr))
        filename = os.path.join('output', str(time.time()) + '.docx')
        print('--- Start generation of doc ', filename, '---')
        doc = Document()
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 0.95
        i = 0
        text = ''
        for item in arr:
            text += item + '                 \t'
            if 3 == i:
                print(text)
                run = p.add_run()
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                run.add_text(text)
                run.add_break()
                run.add_break()
                i = 0
                text = ''
            else:
                i += 1
        doc.save(filename)
        print('--- Complete generation of doc ', filename, '---')

